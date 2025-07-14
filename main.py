import streamlit as st
import os
import re
import uuid
import shutil
import time
import gc
import subprocess
import tempfile
from pathlib import Path
import zipfile
from typing import Dict, List, Any, Optional, Tuple
import json

# Import audio processing dependencies
try:
    from faster_whisper import WhisperModel
    import torch
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

# Import translation dependencies
try:
    from googletrans import Translator
    TRANSLATION_AVAILABLE = True
except ImportError:
    TRANSLATION_AVAILABLE = False

# Import document processing
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Language dictionary from utils.py
LANGUAGE_DICT = {
    "Assamese": {"lang_code": "as", "meta_code": "asm_Beng"},
    "Bengali": {"lang_code": "bn", "meta_code": "ben_Beng"},
    "English": {"lang_code": "en", "meta_code": "eng_Latn"},
    "Gujarati": {"lang_code": "gu", "meta_code": "guj_Gujr"},
    "Hindi": {"lang_code": "hi", "meta_code": "hin_Deva"},
    "Kannada": {"lang_code": "kn", "meta_code": "kan_Knda"},
    "Malayalam": {"lang_code": "ml", "meta_code": "mal_Mlym"},
    "Marathi": {"lang_code": "mr", "meta_code": "mar_Deva"},
    "Punjabi": {"lang_code": "pa", "meta_code": "pan_Guru"},
    "Sinhala": {"lang_code": "si", "meta_code": "sin_Sinh"},
    "Tamil": {"lang_code": "ta", "meta_code": "tam_Taml"},
    "Telugu": {"lang_code": "te", "meta_code": "tel_Telu"},
    "Urdu": {"lang_code": "ur", "meta_code": "urd_Arab"},
}

# Streamlit app configuration
st.set_page_config(
    page_title="Multi-Language Subtitle Generator",
    page_icon="🎬",
    layout="wide"
)

# Global variables
TEMP_DIR = tempfile.mkdtemp()
OUTPUT_DIR = os.path.join(TEMP_DIR, "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Utility functions from your backend
def clean_file_name(file_path):
    """Clean file name for safe processing"""
    file_name = os.path.basename(file_path)
    file_name, file_extension = os.path.splitext(file_name)
    
    # Replace non-alphanumeric characters with underscore
    cleaned = re.sub(r'[^a-zA-Z\d]+', '_', file_name)
    clean_file_name = re.sub(r'_+', '_', cleaned).strip('_')
    
    # Generate random UUID for uniqueness
    random_uuid = uuid.uuid4().hex[:6]
    
    # Combine cleaned file name with extension
    clean_file_path = os.path.join(
        os.path.dirname(file_path), 
        clean_file_name + f"_{random_uuid}" + file_extension
    )
    
    return clean_file_path

def get_language_name(lang_code):
    """Get language name from language code"""
    for language, details in LANGUAGE_DICT.items():
        if details["lang_code"] == lang_code:
            return language
    return lang_code

def format_segments(segments):
    """Format Whisper segments into structured data"""
    saved_segments = list(segments)
    sentence_timestamp = []
    words_timestamp = []
    speech_to_text = ""

    for i in saved_segments:
        # Store sentence information
        text = i.text.strip()
        sentence_id = len(sentence_timestamp)
        sentence_timestamp.append({
            "id": sentence_id,
            "text": text,
            "start": i.start,
            "end": i.end,
            "words": []
        })
        speech_to_text += text + " "

        # Process each word in the sentence
        for word in i.words:
            word_data = {
                "word": word.word.strip(),
                "start": word.start,
                "end": word.end
            }
            sentence_timestamp[sentence_id]["words"].append(word_data)
            words_timestamp.append(word_data)

    return sentence_timestamp, words_timestamp, speech_to_text

def combine_word_segments(words_timestamp, max_words_per_subtitle=8, min_silence_between_words=0.5):
    """Combine words into subtitle segments"""
    if max_words_per_subtitle <= 1:
        max_words_per_subtitle = 1
    
    before_translate = {}
    id = 1
    text = ""
    start = None
    end = None
    word_count = 0
    last_end_time = None

    for i in words_timestamp:
        try:
            word = i['word']
            word_start = i['start']
            word_end = i['end']

            # Check for sentence-ending punctuation
            is_end_of_sentence = word.endswith(('.', '?', '!'))

            # Check for conditions to create a new subtitle
            if ((last_end_time is not None and word_start - last_end_time > min_silence_between_words)
                or word_count >= max_words_per_subtitle
                or is_end_of_sentence):

                # Store the previous subtitle if there's any
                if text:
                    before_translate[id] = {
                        "text": text,
                        "start": start,
                        "end": end
                    }
                    id += 1

                # Reset for the new subtitle segment
                text = word
                start = word_start
                word_count = 1
            else:
                if word_count == 0:
                    start = word_start
                text += " " + word
                word_count += 1

            end = word_end
            last_end_time = word_end

        except KeyError as e:
            print(f"KeyError: {e} - Skipping word")
            pass

    # Add the last subtitle segment
    if text:
        before_translate[id] = {
            "text": text,
            "start": start,
            "end": end
        }

    return before_translate

def convert_time_to_srt_format(seconds):
    """Convert seconds to SRT time format (HH:MM:SS,ms)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    milliseconds = int((seconds - int(seconds)) * 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{milliseconds:03}"

def write_srt_file(subtitles, filename):
    """Write subtitles to SRT file"""
    with open(filename, 'w', encoding='utf-8') as f:
        for id, entry in subtitles.items():
            f.write(f"{id}\n")
            start_time = convert_time_to_srt_format(entry['start'])
            end_time = convert_time_to_srt_format(entry['end'])
            f.write(f"{start_time} --> {end_time}\n")
            f.write(f"{entry['text']}\n\n")

def save_uploaded_file(uploaded_file):
    """Save uploaded file to temporary location"""
    try:
        file_path = os.path.join(TEMP_DIR, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        return file_path
    except Exception as e:
        st.error(f"Error saving file: {e}")
        return None

def transcribe_audio(audio_path, source_language, progress_callback=None):
    """Transcribe audio using Whisper"""
    if not WHISPER_AVAILABLE:
        return None, None, None, "Whisper not available"
    
    try:
        if progress_callback:
            progress_callback(0.1, "🎵 Loading Whisper model...")
        
        # Determine device and compute type
        if torch.cuda.is_available():
            device = "cuda"
            compute_type = "float16"
        else:
            device = "cpu"
            compute_type = "int8"
        
        # Load Whisper model
        model = WhisperModel(
            "deepdml/faster-whisper-large-v3-turbo-ct2",
            device=device,
            compute_type=compute_type
        )
        
        if progress_callback:
            progress_callback(0.3, f"🎙️ Transcribing audio with {device.upper()}...")
        
        # Transcribe audio
        if source_language == "Automatic":
            segments, info = model.transcribe(audio_path, word_timestamps=True)
            detected_language = get_language_name(info.language)
        else:
            lang_code = LANGUAGE_DICT.get(source_language, {}).get("lang_code", "en")
            segments, info = model.transcribe(
                audio_path, 
                word_timestamps=True, 
                language=lang_code
            )
            detected_language = source_language
        
        if progress_callback:
            progress_callback(0.7, "📝 Processing transcription results...")
        
        # Format segments
        sentence_timestamp, words_timestamp, speech_to_text = format_segments(segments)
        
        if progress_callback:
            progress_callback(0.9, "🧹 Cleaning up...")
        
        # Cleanup
        del model
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        
        if progress_callback:
            progress_callback(1.0, f"✅ Transcription complete! Language: {detected_language}")
        
        return sentence_timestamp, words_timestamp, speech_to_text, detected_language
        
    except Exception as e:
        return None, None, None, f"Error transcribing audio: {str(e)}"

def translate_text(text, target_language, progress_callback=None):
    """Translate text using Google Translate"""
    if not TRANSLATION_AVAILABLE:
        return text, "Translation not available"
    
    try:
        translator = Translator()
        target_lang_code = LANGUAGE_DICT.get(target_language, {}).get("lang_code", "en")
        
        if progress_callback:
            progress_callback(0.5, f"🌐 Translating to {target_language}...")
        
        # Split text into chunks to avoid API limits
        chunks = [text[i:i+4500] for i in range(0, len(text), 4500)]
        translated_chunks = []
        
        for i, chunk in enumerate(chunks):
            if progress_callback:
                progress = 0.2 + (0.6 * (i + 1) / len(chunks))
                progress_callback(progress, f"🌐 Translating chunk {i+1}/{len(chunks)}...")
            
            result = translator.translate(chunk, dest=target_lang_code)
            translated_chunks.append(result.text)
            time.sleep(0.1)  # Rate limiting
        
        translated_text = " ".join(translated_chunks)
        
        if progress_callback:
            progress_callback(1.0, f"✅ Translation to {target_language} complete!")
        
        return translated_text, None
        
    except Exception as e:
        return text, f"Translation error: {str(e)}"

def translate_subtitles(subtitles, target_language, progress_callback=None):
    """Translate subtitle segments"""
    if not TRANSLATION_AVAILABLE:
        return subtitles, "Translation not available"
    
    try:
        translator = Translator()
        target_lang_code = LANGUAGE_DICT.get(target_language, {}).get("lang_code", "en")
        translated_subtitles = {}
        
        total_segments = len(subtitles)
        
        for i, (id, entry) in enumerate(subtitles.items()):
            if progress_callback:
                progress = (i + 1) / total_segments
                progress_callback(progress, f"🌐 Translating segment {i+1}/{total_segments}...")
            
            try:
                result = translator.translate(entry['text'], dest=target_lang_code)
                translated_subtitles[id] = {
                    "text": result.text,
                    "start": entry['start'],
                    "end": entry['end']
                }
                time.sleep(0.1)  # Rate limiting
            except Exception as e:
                # Keep original text if translation fails
                translated_subtitles[id] = entry.copy()
        
        return translated_subtitles, None
        
    except Exception as e:
        return subtitles, f"Translation error: {str(e)}"

def create_video_with_subtitles(video_path, srt_path, output_path, progress_callback=None):
    """Create video with embedded subtitles using FFmpeg"""
    try:
        if progress_callback:
            progress_callback(0.1, "🎬 Preparing video processing...")
        
        # FFmpeg command to add subtitles
        cmd = [
            'ffmpeg', '-y',  # -y to overwrite output file
            '-i', video_path,  # Input video
            '-vf', f"subtitles={srt_path}:force_style='FontSize=16,PrimaryColour=&Hffffff,BackColour=&H80000000,Bold=1'",  # Subtitle filter
            '-c:a', 'copy',  # Copy audio stream
            '-preset', 'fast',  # Fast encoding preset
            output_path
        ]
        
        if progress_callback:
            progress_callback(0.3, "🎬 Processing video with subtitles...")
        
        # Run FFmpeg command
        process = subprocess.run(cmd, capture_output=True, text=True)
        
        if process.returncode == 0:
            if progress_callback:
                progress_callback(1.0, "✅ Video with subtitles created successfully!")
            return True, None
        else:
            error_msg = f"FFmpeg error: {process.stderr}"
            return False, error_msg
            
    except Exception as e:
        return False, f"Error creating video: {str(e)}"

def create_subtitle_document(subtitles, filename, language):
    """Create a DOCX document with subtitles"""
    if not DOCX_AVAILABLE:
        return False, "python-docx not available"
    
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Subtitles - {language}', 0)
        
        # Add metadata
        doc.add_paragraph(f"Language: {language}")
        doc.add_paragraph(f"Total Segments: {len(subtitles)}")
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()  # Empty line
        
        # Add subtitles
        for id, entry in subtitles.items():
            # Add segment number and timing
            p = doc.add_paragraph()
            p.add_run(f"Segment {id}").bold = True
            p.add_run(f" ({convert_time_to_srt_format(entry['start'])} --> {convert_time_to_srt_format(entry['end'])})")
            
            # Add subtitle text
            doc.add_paragraph(entry['text'])
            doc.add_paragraph()  # Empty line between segments
        
        # Save document
        doc.save(filename)
        return True, None
        
    except Exception as e:
        return False, f"Error creating document: {str(e)}"

def check_ffmpeg():
    """Check if FFmpeg is available"""
    try:
        subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False

def main():
    """Main Streamlit application"""
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #ff6b6b, #4ecdc4);
        padding: 1.5rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .status-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #007bff;
        margin: 1rem 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🎬 Multi-Language Subtitle Generator</h1>
        <p>Convert speech to subtitles, translate to multiple languages, and create videos with embedded subtitles</p>
        <p style="font-size: 0.9em; opacity: 0.9;">✅ Speech Recognition • 🌐 Multi-Language Translation • 🎬 Video Processing • 📄 Document Export</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar with system status
    with st.sidebar:
        st.header("🔧 System Status")
        
        # Check dependencies
        st.markdown("### 📦 Dependencies")
        if WHISPER_AVAILABLE:
            st.success("✅ Faster Whisper")
        else:
            st.error("❌ Faster Whisper")
        
        if TRANSLATION_AVAILABLE:
            st.success("✅ Google Translate")
        else:
            st.error("❌ Google Translate")
        
        if DOCX_AVAILABLE:
            st.success("✅ Document Processing")
        else:
            st.error("❌ Document Processing")
        
        if check_ffmpeg():
            st.success("✅ FFmpeg")
        else:
            st.error("❌ FFmpeg")
        
        # GPU status
        st.markdown("### 🎮 Hardware")
        if torch.cuda.is_available():
            st.success("✅ GPU Available")
        else:
            st.info("💻 CPU Only")
        
        st.divider()
        
        # Language selection
        st.header("🌐 Languages")
        available_languages = list(LANGUAGE_DICT.keys())
        st.write(f"**Available:** {len(available_languages)} languages")
        
        with st.expander("View All Languages"):
            for lang, details in LANGUAGE_DICT.items():
                st.write(f"🌍 {lang} ({details['lang_code']})")
    
    # Main interface
    st.header("📤 Upload Audio/Video File")
    
    uploaded_file = st.file_uploader(
        "Choose an audio or video file",
        type=['mp3', 'wav', 'mp4', 'avi', 'mov', 'mkv', 'flv', 'm4a', 'aac'],
        help="Upload audio/video files for subtitle generation"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        # File info
        file_size = len(uploaded_file.getbuffer())
        col1, col2 = st.columns(2)
        with col1:
            st.metric("File Size", f"{file_size/1024/1024:.1f} MB")
        with col2:
            st.metric("File Type", uploaded_file.name.split('.')[-1].upper())
        
        # Configuration
        st.header("⚙️ Configuration")
        
        col1, col2 = st.columns(2)
        with col1:
            # Source language selection
            source_language = st.selectbox(
                "Source Language",
                ["Automatic"] + available_languages,
                help="Select the language spoken in the audio"
            )
        
        with col2:
            # Max words per subtitle
            max_words = st.slider(
                "Max Words per Subtitle",
                min_value=1,
                max_value=15,
                value=8,
                help="Maximum number of words per subtitle segment"
            )
        
        # Target languages for translation
        st.subheader("🌐 Translation Languages")
        target_languages = st.multiselect(
            "Select languages for translation",
            available_languages,
            default=["English", "Hindi", "Bengali"],
            help="Choose which languages to translate subtitles into"
        )
        
        # Processing options
        st.subheader("📋 Output Options")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            generate_srt = st.checkbox("Generate SRT files", value=True)
        with col2:
            generate_video = st.checkbox("Generate videos with subtitles", value=True)
        with col3:
            generate_docs = st.checkbox("Generate DOCX documents", value=True)
        
        # Process button
        if st.button("🚀 Start Processing", type="primary"):
            if not any([generate_srt, generate_video, generate_docs]):
                st.error("❌ Please select at least one output option")
                return
            
            # Save uploaded file
            input_file_path = save_uploaded_file(uploaded_file)
            if not input_file_path:
                st.error("❌ Failed to save uploaded file")
                return
            
            # Initialize progress tracking
            overall_progress = st.progress(0)
            status_text = st.empty()
            
            try:
                # Step 1: Transcribe audio
                status_text.text("🎙️ Transcribing audio...")
                overall_progress.progress(0.1)
                
                transcribe_progress = st.progress(0)
                transcribe_status = st.empty()
                
                def update_transcribe_progress(value, message):
                    transcribe_progress.progress(value)
                    transcribe_status.text(message)
                
                sentence_timestamp, words_timestamp, speech_to_text, detected_language = transcribe_audio(
                    input_file_path, source_language, update_transcribe_progress
                )
                
                if sentence_timestamp is None:
                    st.error(f"❌ Transcription failed: {detected_language}")
                    return
                
                st.success(f"✅ Transcription complete! Detected language: {detected_language}")
                
                # Clear transcription progress
                transcribe_progress.empty()
                transcribe_status.empty()
                
                # Create subtitle segments
                status_text.text("📝 Creating subtitle segments...")
                overall_progress.progress(0.2)
                
                word_segments = combine_word_segments(words_timestamp, max_words, 0.5)
                
                # Step 2: Generate original language outputs
                status_text.text("📄 Generating original language files...")
                overall_progress.progress(0.3)
                
                base_filename = os.path.splitext(uploaded_file.name)[0]
                results = {"original": {}, "translations": {}}
                
                # Original language SRT
                if generate_srt:
                    original_srt_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{detected_language}.srt")
                    write_srt_file(word_segments, original_srt_path)
                    results["original"]["srt"] = original_srt_path
                
                # Original language DOCX
                if generate_docs:
                    original_doc_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{detected_language}.docx")
                    success, error = create_subtitle_document(word_segments, original_doc_path, detected_language)
                    if success:
                        results["original"]["doc"] = original_doc_path
                
                # Step 3: Translate to target languages
                if target_languages and detected_language not in target_languages:
                    target_languages.append(detected_language)  # Ensure original language is included
                
                translation_progress_step = 0.4 / len(target_languages) if target_languages else 0
                
                for i, target_lang in enumerate(target_languages):
                    if target_lang == detected_language:
                        continue  # Skip original language
                    
                    status_text.text(f"🌐 Translating to {target_lang}...")
                    overall_progress.progress(0.3 + (i + 1) * translation_progress_step)
                    
                    # Create translation progress
                    trans_progress = st.progress(0)
                    trans_status = st.empty()
                    
                    def update_trans_progress(value, message):
                        trans_progress.progress(value)
                        trans_status.text(message)
                    
                    # Translate subtitles
                    translated_subtitles, error = translate_subtitles(
                        word_segments, target_lang, update_trans_progress
                    )
                    
                    if error:
                        st.warning(f"⚠️ Translation to {target_lang} failed: {error}")
                        continue
                    
                    results["translations"][target_lang] = {}
                    
                    # Generate translated SRT
                    if generate_srt:
                        trans_srt_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{target_lang}.srt")
                        write_srt_file(translated_subtitles, trans_srt_path)
                        results["translations"][target_lang]["srt"] = trans_srt_path
                    
                    # Generate translated DOCX
                    if generate_docs:
                        trans_doc_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{target_lang}.docx")
                        success, error = create_subtitle_document(translated_subtitles, trans_doc_path, target_lang)
                        if success:
                            results["translations"][target_lang]["doc"] = trans_doc_path
                    
                    # Clear translation progress
                    trans_progress.empty()
                    trans_status.empty()
                
                # Step 4: Generate videos with subtitles
                if generate_video and uploaded_file.name.split('.')[-1].lower() in ['mp4', 'avi', 'mov', 'mkv']:
                    status_text.text("🎬 Creating videos with subtitles...")
                    overall_progress.progress(0.7)
                    
                    # Generate video for original language
                    if "srt" in results["original"]:
                        video_progress = st.progress(0)
                        video_status = st.empty()
                        
                        def update_video_progress(value, message):
                            video_progress.progress(value)
                            video_status.text(message)
                        
                        output_video_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{detected_language}_subtitled.mp4")
                        success, error = create_video_with_subtitles(
                            input_file_path, results["original"]["srt"], output_video_path, update_video_progress
                        )
                        
                        if success:
                            results["original"]["video"] = output_video_path
                        else:
                            st.warning(f"⚠️ Video creation failed: {error}")
                        
                        video_progress.empty()
                        video_status.empty()
                    
                    # Generate videos for translated languages
                    for lang, files in results["translations"].items():
                        if "srt" in files:
                            video_progress = st.progress(0)
                            video_status = st.empty()
                            
                            output_video_path = os.path.join(OUTPUT_DIR, f"{base_filename}_{lang}_subtitled.mp4")
                            success, error = create_video_with_subtitles(
                                input_file_path, files["srt"], output_video_path, update_video_progress
                            )
                            
                            if success:
                                results["translations"][lang]["video"] = output_video_path
                            
                            video_progress.empty()
                            video_status.empty()
                
                # Step 5: Create download package
                status_text.text("📦 Creating download package...")
                overall_progress.progress(0.9)
                
                zip_path = os.path.join(OUTPUT_DIR, f"{base_filename}_complete_package.zip")
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    # Add original files
                    for file_type, file_path in results["original"].items():
                        if os.path.exists(file_path):
                            zipf.write(file_path, f"original/{os.path.basename(file_path)}")
                    
                    # Add translated files
                    for lang, files in results["translations"].items():
                        for file_type, file_path in files.items():
                            if os.path.exists(file_path):
                                zipf.write(file_path, f"translations/{lang}/{os.path.basename(file_path)}")
                
                # Complete
                overall_progress.progress(1.0)
                status_text.text("✅ Processing complete!")
                
                # Display results
                st.header("📥 Download Results")
                
                # Summary
                col1, col2, col3 = st.columns(3)
                with col1:
                    total_files = len(results["original"]) + sum(len(files) for files in results["translations"].values())
                    st.metric("Total Files Generated", total_files)
                with col2:
                    st.metric("Languages Processed", 1 + len(results["translations"]))
                with col3:
                    st.metric("Original Language", detected_language)
                
                # Download complete package
                st.subheader("📦 Complete Package")
                if os.path.exists(zip_path):
                    with open(zip_path, "rb") as f:
                        st.download_button(
                            label="📥 Download Complete Package (ZIP)",
                            data=f.read(),
                            file_name=f"{base_filename}_complete_package.zip",
                            mime="application/zip"
                        )
                
                # Individual file downloads
                st.subheader("📄 Individual Files")
                
                # Original language files
                st.write(f"**Original Language: {detected_language}**")
                col1, col2, col3 = st.columns(3)
                
                if "srt" in results["original"]:
                    with col1:
                        with open(results["original"]["srt"], "rb") as f:
                            st.download_button(
                                f"📄 {detected_language} SRT",
                                data=f.read(),
                                file_name=os.path.basename(results["original"]["srt"]),
                                mime="text/plain"
                            )
                
                if "doc" in results["original"]:
                    with col2:
                        with open(results["original"]["doc"], "rb") as f:
                            st.download_button(
                                f"📄 {detected_language} DOCX",
                                data=f.read(),
                                file_name=os.path.basename(results["original"]["doc"]),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                
                if "video" in results["original"]:
                    with col3:
                        with open(results["original"]["video"], "rb") as f:
                            st.download_button(
                                f"🎬 {detected_language} Video",
                                data=f.read(),
                                file_name=os.path.basename(results["original"]["video"]),
                                mime="video/mp4"
                            )
                
                # Translated files
                if results["translations"]:
                    st.write("**Translated Languages:**")
                    for lang, files in results["translations"].items():
                        st.write(f"**{lang}:**")
                        col1, col2, col3 = st.columns(3)
                        
                        if "srt" in files:
                            with col1:
                                with open(files["srt"], "rb") as f:
                                    st.download_button(
                                        f"📄 {lang} SRT",
                                        data=f.read(),
                                        file_name=os.path.basename(files["srt"]),
                                        mime="text/plain",
                                        key=f"srt_{lang}"
                                    )
                        
                        if "doc" in files:
                            with col2:
                                with open(files["doc"], "rb") as f:
                                    st.download_button(
                                        f"📄 {lang} DOCX",
                                        data=f.read(),
                                        file_name=os.path.basename(files["doc"]),
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"doc_{lang}"
                                    )
                        
                        if "video" in files:
                            with col3:
                                with open(files["video"], "rb") as f:
                                    st.download_button(
                                        f"🎬 {lang} Video",
                                        data=f.read(),
                                        file_name=os.path.basename(files["video"]),
                                        mime="video/mp4",
                                        key=f"video_{lang}"
                                    )
                
                st.success("🎉 All files generated successfully!")
                
            except Exception as e:
                st.error(f"❌ Processing failed: {str(e)}")
                
            finally:
                # Cleanup
                if os.path.exists(input_file_path):
                    os.remove(input_file_path)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.9em;'>
        <p>🎬 Multi-Language Subtitle Generator | Speech Recognition • Translation • Video Processing</p>
        <p>🔒 Secure processing • 🌐 13 Indian languages • 🎥 Video subtitle embedding • 📄 Document export</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
